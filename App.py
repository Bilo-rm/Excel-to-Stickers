import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_table_alignment(table):
    """Center the table on the page."""
    tbl = table._element
    tbl_pr = tbl.find(qn('w:tblPr'))  # Get table properties
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.append(tbl_pr)

    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')  # Center the table
    tbl_pr.append(jc)

def create_sticker_doc(excel_file, output_docx):
    df = pd.read_excel(excel_file)
    doc = Document()
    
    index = 0
    for _, row in df.iterrows():
        if index % 16 == 0:  
            if index > 0:
                doc.add_page_break()
            table = doc.add_table(rows=8, cols=2)
            table.autofit = False
            set_table_alignment(table)  # Center the table

            for r in table.rows:
                for cell in r.cells:
                    cell.width = Cm(10.5)
                    cell.height = Cm(3.7125)

                    # Center the content vertically
                    tc_pr = cell._element.get_or_add_tcPr()
                    v_align = OxmlElement('w:vAlign')
                    v_align.set(qn('w:val'), 'center')
                    tc_pr.append(v_align)

        id_value = row['ID']
        full_name = f"{row['Fname']} {row['Lname']}"
        department = row['Department']
        
        # Insert data into the correct cell
        cell = table.cell((index % 16) // 2, (index % 16) % 2)
        cell.text = f"{id_value}\n{full_name}\n{department}"

        # Format text
        for paragraph in cell.paragraphs:
            paragraph.alignment = 1  # Center horizontally
            run = paragraph.runs[0]
            run.font.size = Pt(20)  

        index += 1

    # Ensure file is not open before saving
    try:
        doc.save(output_docx)
        print(f"Document saved as {output_docx}")
    except PermissionError:
        print(f"Error: Close '{output_docx}' before running the script again.")

create_sticker_doc("Test.xlsx", "stickers6.docx")
